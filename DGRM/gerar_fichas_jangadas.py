# Script para gerar ficha de identificação de jangada em Word para cada jangada cadastrada
# Usa python-docx
import os
from docx import Document
from datetime import datetime

# Exemplo de dados de jangadas
jangadas = [
    {
        'numero_serie': 'RFD12345',
        'marca': 'RFD',
        'modelo': 'MKIV',
        'capacidade': '10',
        'ano_fabricacao': '2022',
        'pais_fabricacao': 'Inglaterra',
        'cliente': 'Navio X',
        'data_inspecao': '07/02/2026',
    },
    {
        'numero_serie': 'DSB67890',
        'marca': 'DSB',
        'modelo': 'LR05',
        'capacidade': '12',
        'ano_fabricacao': '2021',
        'pais_fabricacao': 'Alemanha',
        'cliente': 'Navio Y',
        'data_inspecao': '07/02/2026',
    },
    {
        'numero_serie': 'ZOD54321',
        'marca': 'ZODIAC',
        'modelo': 'ZHR',
        'capacidade': '8',
        'ano_fabricacao': '2020',
        'pais_fabricacao': 'França',
        'cliente': 'Navio Z',
        'data_inspecao': '07/02/2026',
    },
]

def gerar_ficha(jangada):
    doc = Document()
    doc.add_heading('Ficha de Identificação de Jangada', 0)
    doc.add_paragraph(f"Número de Série: {jangada['numero_serie']}")
    doc.add_paragraph(f"Marca: {jangada['marca']}")
    doc.add_paragraph(f"Modelo: {jangada['modelo']}")
    doc.add_paragraph(f"Capacidade: {jangada['capacidade']}")
    doc.add_paragraph(f"Ano de Fabricação: {jangada['ano_fabricacao']}")
    doc.add_paragraph(f"País de Fabricação: {jangada['pais_fabricacao']}")
    doc.add_paragraph(f"Cliente: {jangada['cliente']}")
    doc.add_paragraph(f"Data da Inspeção: {jangada['data_inspecao']}")
    nome_arquivo = f"DGRM/ficha de identificação de jangada - {jangada['numero_serie']}.docx"
    doc.save(nome_arquivo)
    print(f"Gerado: {nome_arquivo}")

if not os.path.exists('DGRM'):
    os.makedirs('DGRM')

for jangada in jangadas:
    gerar_ficha(jangada)
